import os
import tkinter
import whisper
from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk
from docx import Document
from docx2pdf import convert



model = whisper.load_model("base")
filepath = ""
filedest = ""



def openDestination():
    global filedest
    filedest = filedialog.askdirectory()
    templabel = tkinter.Label(destFrame, text=filedest, font=("Arial, 8"), bg="#273746", fg="white", padx="20").grid(column=1, row=6, sticky=(N, S))


def openDirectory():
    global filepath
    filepath = filedialog.askopenfilename()
    templabel = tkinter.Label(dirFrame, text=filepath, font=("Arial, 8"), bg="#273746", fg="white", padx="20").grid(column=1, row=4, sticky=(N, S))
    return filepath

def transcribe_audio(audio_path):
    result = model.transcribe(audio_path)
    return result['text']
def transcribe_text():
    global filepath
    print("This is the filepath: " + filepath)
    audio_path = filepath
    transcribed_text = transcribe_audio(audio_path)
    file = open(filedest + "/transcription.txt", "w")
    file.write(transcribed_text)
    messagebox.showinfo("Success!",message="Transcription succeeded!")

def transcribe_docx():
    global filepath
    audio_path = filepath
    transcribed_text = transcribe_audio(audio_path)
    document = Document()
    document.add_paragraph(transcribed_text)
    document.save(filedest + "/transcription.docx")
    messagebox.showinfo("Success!", message="Transcription succeeded!")

def transcribe_pdf():
    global filepath
    audio_path = filepath
    transcribed_text = transcribe_audio(audio_path)
    document = Document()
    document.add_paragraph(transcribed_text)
    document.save(filedest + "/transcription.docx")
    convert(filedest + "/transcription.docx", filedest + "/transcription.pdf")
    os.remove(filedest + "/transcription.docx")
    messagebox.showinfo("Success!", message="Transcription succeeded!")




window = Tk()
window.title("Transcriptor")
window.configure(bg="#273746", padx="20", pady="20")
window.columnconfigure(0, weight=1)
window.rowconfigure(0, weight=1)

dirFrame = tkinter.Frame()
destFrame = tkinter.Frame()
transFrame = tkinter.Frame()

label = tkinter.Label(window, text="AI Transcription Tool", font=("Arial, 25"), bg="#273746", fg="white")
button = tkinter.Button(dirFrame, text="Choose directory", padx="20", font=("Arial, 8"), bg="#273746", fg="white", command=openDirectory)
dir_text = tkinter.Label(dirFrame, text="Chose the directory of your audio file: ", font=("Arial, 8"), bg="#273746", fg="white")
dirFrame.configure(background="#273746", padx="10", pady="10", highlightthickness="5", highlightbackground="white")

destination = tkinter.Label(destFrame, text="Chose the transcription destination: ", font=("Arial, 8"), bg="#273746", fg="white")
destbutton = tkinter.Button(destFrame, text="Destination", font=("Arial, 8"), bg="#273746", fg="white", padx="20", command=openDestination)
destFrame.configure(background="#273746", padx="10", pady="10", highlightthickness="5", highlightbackground="white")

transcribetext_button = tkinter.Button(transFrame, text="Transcribe to Text", font=("Arial, 8"), bg="#273746", fg="white", command=transcribe_text)
transcribedocx_button = tkinter.Button(transFrame, text="Transcribe to DOCX", font=("Arial, 8"), bg="#273746", fg="white", command=transcribe_docx)
transcribepdf_button = tkinter.Button(transFrame, text="Transcribe to PDF", font=("Arial, 8"), bg="#273746", fg="white", command=transcribe_pdf)
transFrame.configure(background="#273746", padx="10", pady="10")

#Grid manipulation
label.grid(column=1, row=1, rowspan=2, pady=8)

dirFrame.grid(column=1, row=3, pady=5)
dir_text.grid(column=1, row=3, padx=5)
button.grid(column=2, row=3, padx=5)

destFrame.grid(column=1, row=5, pady=5)
destination.grid(column=1, row=5, padx=5)
destbutton.grid(column=2, row=5, padx=5)

transFrame.grid(column=1, row=7)
transcribetext_button.grid(column=1, row=7, padx=5)
transcribedocx_button.grid(column=2, row=7, padx=5)
transcribepdf_button.grid(column=3, row=7, padx=5)
window.mainloop()